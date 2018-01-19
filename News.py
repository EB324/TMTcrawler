import requests
from bs4 import BeautifulSoup
from docx import Document
import time
import re

#document = Document()
#document.add_heading('行业新闻')

def getNews(url):
    html = requests.get(url)
    soup = BeautifulSoup(html.text, 'html.parser')
    newsList = soup.select('.newslist a')

    for news in newsList:
        html2 = requests.get(news.get('href'))
        soup2 = BeautifulSoup(html2.text, 'html.parser')

        if soup2.select('.text'):
            #            p = document.add_paragraph()
            #            p.add_run(news.get_text()).bold = True
            print(news.get_text())
            print(news.get('href'))

            newsDetail = soup2.select('.text')
            i = 0
            while len(newsDetail[i].text) <= 30:
                i = i+1
            else:
#                p = document.add_paragraph(newsDetail[i].text)
#                p.add_run('\n')
                print(newsDetail[i].text, '\n')

time = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())
timeSplit = re.split('-| ', time)

year = timeSplit[0]
month = timeSplit[1]
date = timeSplit[2]

newsUrl = 'http://tech.qq.com/l/'+year+month+'/scroll_'+date+'.htm'
getNews(newsUrl)

# 如果有第二面
newsUrl2 = newsUrl.rstrip('.htm') + '_2' + '.htm'
if (newsUrl2):
    getNews(newsUrl2)

#document.save('News.docx')